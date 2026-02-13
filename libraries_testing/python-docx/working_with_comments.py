from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


# -----------------------------
# Step 1: Create a document
# -----------------------------
document = Document()

# -----------------------------
# Step 2: Add a paragraph
# -----------------------------
paragraph_1 = document.add_paragraph("Hello, World")
paragraph_2 = document.add_paragraph("The Rain in spain")
paragraph_3 = document.add_paragraph("India is the Poluted contry in World")

run_special = paragraph_2.add_run("falls mainly in the plain.")

# ----------------------------------------
# Step 3: Add a SIMPLE comment to a whole paragraph
# (Comment Reference = paragraph.runs)
# ----------------------------------------
comment_1 = document.add_comment(
    runs=paragraph_2.runs,                  # range = full paragraph
    text="I have this to say about that.",
    author="Steve Canny",
    initials="SC"
)

# -----------------------------------------------
# Step 4: Add a comment to a SPECIFIC RANGE (run)
# (Comment Reference = selected text only)
# -----------------------------------------------
comment_2 = document.add_comment(
    runs=[run_special],
    text="Is this phrase still accurate?",
    author="Reviewer",
    initials="RV"
)

# -------------------------------------------------
# Step 5: Add RICH CONTENT inside a comment
# (Comment Content = block-item container)
# -------------------------------------------------
comment_3 = document.add_comment(
    runs=paragraph_2.runs,
    text="",                      # empty initial content
    author="Editor",
    initials="ED",
)

# Access first paragraph inside comment-content
c_para = comment_3.paragraphs[0]

c_para.add_run("Please finish this thought. I believe it should be ")
c_para.add_run("falls mainly in the plain.").bold = True
c_para.add_run(" (classic example)").italic = True

# ------------------------------------------------
# Step 6: Update COMMENT METADATA
# ------------------------------------------------
comment_3.author = "John Smith"
comment_3.initials = "JS"

# -----------------------------------------
# Step 7: Access COMMENT COLLECTION
# -----------------------------------------
comments = document.comments

print("Total Comments:", len(comments))

for c in comments:
    print("ID:", c.id)
    print("Author:", c.author)
    print("Initials:", c.initials)
    print("Text:", c.text)
    print("-" * 40)

# -----------------------------
# Step 8: Save the document
# -----------------------------
document.save("./ocr_data_doc/working_with_comments.docx")
print("Programm RUN Successfully..! | 'working_with_comments.docx'")


