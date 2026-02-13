from docx import Document

document = Document()

# ---------------------------------------------------
# ------- Accessing the header for a section --------
# ---------------------------------------------------


# ----- Step 1: Add content to the document body ------
document.add_paragraph("This is the main body of the page 1.")
document.add_page_break()
document.add_paragraph("This is the main body of the page 2.")


# ------ Step 2: Access the first section -------
section = document.sections[0]


# ------ Step 3: Add a header --------
header = section.header
header_para = header.paragraphs[0]
header_para.text = "My Report Header - Page Number:  \t"


# ------ Step 4: Add a Footer -------
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = "Confidectial - Do Not Share"



document.save("./ocr_data_doc/working_with_headers_and_footers.docx")
print("Programm RUN SUCCESS, it was run properlly OK!")
