from docx import Document

# 1. Create a document (main API object)
document = Document()

# 2. Add a heading (it's just a paragraph with a style)
document.add_heading("Python-docx API Basics", level=1)

# 3. Add a normal paragraph
paragraph = document.add_paragraph("This document is created using ")

# 4. Inline object: add text inside the pagragraph
run = paragraph.add_run("Python-docx")
run.bold = True

paragraph.add_run(" Library.")

# 5. Add bullet list (still paragraph, just styled)
document.add_paragraph("Easy to use", style="List Bullet")
document.add_paragraph("Write ducument top to bottom", style="List Bullet")
document.add_paragraph("Powerfull when needed", style="List Bullet")

# 5. Add a table (block-level object)
table = document.add_table(rows=2, cols=2)
table.style = "Table Grid"

table.cell(0, 0).text = "Feature"
table.cell(0, 1).text = "Description"
table.cell(1, 0).text = "API"
table.cell(1, 1).text = "Simple but Powerful"

# 7. Save the document
document.save("./ocr_data_doc/API_Basics.docx")
