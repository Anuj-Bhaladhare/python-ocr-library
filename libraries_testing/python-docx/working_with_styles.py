from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches

document = Document()
styles = document.styles


# --------------------------------------------------
# Step 2: Use an existing built-in style (Heading 1)
# --------------------------------------------------
tital_1 = document.add_paragraph("This is Tital Number 1111 ")
tital_2 = document.add_paragraph("This is Tital Number 2222 ")
tital_3 = document.add_paragraph("This is Tital Number 3333 ")
tital_4 = document.add_paragraph("This is Tital Number 4444 ")

tital_1.style = styles['Heading 1']
tital_2.style = styles['Heading 2']
tital_3.style = styles['Heading 3']
tital_4.style = styles['Heading 4']

document.add_page_break()


# ------------------------------------
# Step 3: Add body text (Normal style)
# ------------------------------------
paragraph_1 = document.add_paragraph("This report shows the sales preformance for the current month.")
paragraph_1.style = styles['Normal']
document.add_page_break()


# --------------------------------------
# Step 4: Apply style by name (shortcut)
# --------------------------------------
document.add_paragraph(
    "Key observations are listed below:",
    style = "Body Text"
)
document.add_page_break()


# ---------------------------------------
# Step 5: Use a list style || Bullet list
# ---------------------------------------
document.add_paragraph("Sales increased by 12%", style="List Bullet")
document.add_paragraph("New customers grew by 8%", style="List Bullet")
document.add_paragraph("Returns reduce 13%", style="List Bullet")
document.add_page_break()

# -----------------------------------------------------
# Step 6: Create your OWN custom style (very important)
# -----------------------------------------------------
note_style = styles.add_style("NoteText", WD_STYLE_TYPE.PARAGRAPH)
document.add_page_break()


# ----------------------------------------
# Step 7: Inherit from Normal (base_style)
# ----------------------------------------
note_style.base_style = styles['Normal']
document.add_page_break()


# ------------------------------------------
# Step 8: Define character formatting (font)
# ------------------------------------------
font = note_style.font
font.italic = True
font.size = Pt(10)
document.add_page_break()


# ------------------------------------------------------
# Step 9: Define paragraph formatting (indent + spacing)
# ------------------------------------------------------
pf = note_style.paragraph_format
pf.left_indent = Inches(0.5)
pf.space_before = Pt(6)
document.add_page_break()


# -----------------------------
# Step 10: Use the custom style
# -----------------------------
document.add_paragraph(
    "Note: Data is provisional and subject to revision.",
    style="NoteText"
)
document.add_page_break()


# ---------------------------------------------------------
# Step 11: Control next paragraph behavior (Heading → Body)
# ---------------------------------------------------------
styles['Heading 1'].next_paragraph_style = styles['Body Text']
document.add_page_break()


# -----------------------------------------
# Step 12: Show style in Word Style Gallery
# -----------------------------------------
note_style.hidden = False
note_style.quick_style = True
note_style.priority = 5
document.add_page_break()


# --------------------------
# Step 13: Save the document
# --------------------------
document.save("./ocr_data_doc/working_with_styles.docx")
print("Document save properlly...! || OK!")

