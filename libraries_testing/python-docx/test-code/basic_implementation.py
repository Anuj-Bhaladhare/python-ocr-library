from docx import Document
from docx.shared import Inches

def create_word_from_ocr(ocr_data, output_path="output.docx"):
    """
    Create a Word document from OCR results
    
    Args:
        ocr_data: List of dictionaries containing OCR text and metadata
        output_path: Path to save the .docx file
    """
    # Create a new document
    doc = Document()
    
    # Add a title
    doc.add_heading("This is the headuibg", 0)
    
    # Process each OCR result
    for i, ocr_item in enumerate(ocr_data):
        # Add heading for each page/section
        doc.add_heading(f'Page {i+1}', level=1)
        
        # Add the OCR text
        paragraph = doc.add_paragraph(ocr_item['text'])
        
        # Add metadata if available
        if 'confidence' in ocr_item:
            doc.add_paragraph(f'Confidence: {ocr_item["confidence"]:.2f}%')
        
        # Add page break between pages (except after last page)
        if i < len(ocr_data) - 1:
            doc.add_page_break()
    
    # Save the document
    doc.save(output_path)
    print(f"Document saved to {output_path}")

