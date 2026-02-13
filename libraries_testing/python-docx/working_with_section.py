from docx import Document
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from docx.shared import Inches

document = Document()

# -----------------------------
# ---- Accessing sections -----
# -----------------------------

# ------ Section 1 (Default) -------
paragraph_1 = document.add_paragraph("Section 1: Portrait Page")
section_1 = document.sections[0]

section_1.top_margin = Inches(1)     # Aided top marging in section
section_1.bottom_margin = Inches(1)  # Aided bottom marging in section

# ------ Section 2 (Landscape) ------
section_2 = document.add_section(WD_SECTION_START.NEW_PAGE)        # Adding a new section
paragraph_2 = document.add_paragraph("Section 2: Landscape page")

# Orientation in Landscape of the page
section_2.orientation = WD_ORIENTATION.LANDSCAPE
section_2.page_width, section_2.page_height = (
    section_2.page_height,
    section_2.page_width
)

# ------ Section 3 (Odd page start) ------
section_3 = document.add_section(WD_SECTION_START.ODD_PAGE)        # Adding a new section
paragraph_3 = document.add_paragraph("Section 3: Starts on odd page")

# Orientation in portrate of the page
section_3.orientation = WD_ORIENTATION.PORTRAIT
section_3.page_width, section_3.page_height = (
    section_3.page_height,
    section_3.page_width
)



document.save("./ocr_data_doc/working_with_section.docx")
print("Programm Run Successfully...")
