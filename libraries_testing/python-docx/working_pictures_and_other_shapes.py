from docx import Document
from docx.shared import Inches

document = Document()

# --------------------------------------------------
# 🧪 Example 3: Inline Image Resize (Very Important)
# --------------------------------------------------
document.add_paragraph("Company Logo:")
document.add_picture("./images/logo-logo.png", width=Inches(1.5))
document.add_page_break()


# -----------------------------------------------------------
# 🧪 Example 4: Proof that image is INLINE (Text delete test)
# -----------------------------------------------------------
document.add_heading("Inspection Report", level=1)
document.add_paragraph("Below image is captured from site inspection:")
document.add_picture("./images/nature.jpg")
document.add_paragraph("Observation: Minor crack found near the pillar.")
document.add_page_break()


# -------------------------------------------------------------
# 🧪 Example 5: Multiple Images in Loop (OCR / Automation Case)
# -------------------------------------------------------------
images = [
    "./images/test_image_1.jpeg", 
    "./images/test_image_2.jpeg", 
    "./images/test_image_3.jpeg"
]
document.add_heading("OCR Captured Images", level=1)

for img in images:
    document.add_paragraph("Captured Image:")
    document.add_picture(img, width=Inches(3))

    

document.save("./ocr_data_doc/working_pictures_and_other_shapes.docx")
print("Programm RUN Success || OK..! Working..!")
