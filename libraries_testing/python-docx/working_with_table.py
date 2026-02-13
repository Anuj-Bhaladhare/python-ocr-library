from docx import Document

document = Document()

# Create a table
table = document.add_table(cols=4, rows=4)

# Access a first row
first_row = table.rows[0]
second_row = table.rows[1]
fourth_row = table.rows[3]


# Murge the table
cell_15 = table.cell(3, 2)
cell_16 = table.cell(3, 3)
cell_15.merge(cell_16)
cell_15.text = "murged 15, 16"

# Murge the table
cell_6 = table.cell(1, 1)
cell_10 = table.cell(2, 1)
cell_6.merge(cell_10)
cell_6.text = "murged 6, 10"

# ----------- Complication 2: Omitted Cells -------------
# Row 1: 3 cells
table = document.add_table(rows=1, cols=3)
row1 = table.rows[0].cells
row1[0].text = "a"
row1[1].text = "b"
row1[2].text = "c"

# Row 2: OMIT first cell (only 2 cells)
row2 = table.add_row().cells
row2[0].text = "d"
row2[1].text = "e"
# <-- no third cell --> omitted

# Row 3: 3 cell again
row3 = table.add_row().cells
row3[0].text = "f"
row3[1].text = "g"
row3[2].text = "h"

document.add_page_break()

table2 = document.add_table(rows=1, cols=2)

# Header row
hdr = table2.rows[0].cells
hdr[0].text = "A"
hdr[1].text = "B"

# Row A
tab2_row1 = table2.add_row().cells
tab2_row1[0].text = "1"
tab2_row1[1].text = "0"

# Row B
tab2_row2 = table2.add_row().cells
tab2_row2[0].text = "0"
tab2_row2[1].text = "1"

document.add_page_break()

"""
    +---+---+   +---+---+---+
    | a | b |   | a | b | c |
+---+---+---+   +---+---+---+
| c | d |           | d |
+---+---+       +---+---+---+
    | e |       | e | f | g |
    +---+       +---+---+---+
"""
# Step 1: base grid = 3 columns
table3 = document.add_table(rows=0, cols=3)


# -------------------------
# Row 1: 1 Omitted cell at START -> [a, b]
# -------------------------
tab3_row1 = table3.add_row().cells
tab3_row1[0].text = "a"
tab3_row1[1].text = "b"
# first grid position is implicitly omitted


# -------------------------
# Row 2: 1 Omitted cell at END -> [c, d]
# -------------------------
tab3_row2 = table3.add_row().cells
tab3_row2[0].text = "c"
tab3_row2[1].text = "d"
# third cell implicitly omitted at end


# -------------------------
# Row 3: 2 Omitted cells (START + END) -> [e]
# -------------------------
tab3_row3 = table3.add_row().cells
tab3_row3[0].text = "e"
# one cell at start and one at end implicitly omitted



document.save("./ocr_data_doc/working_with_table.docx")
print("Run Success...!")

