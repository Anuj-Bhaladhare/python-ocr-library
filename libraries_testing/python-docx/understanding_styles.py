from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


# Create a Document Object OR define document
document = Document()


# ---------------------------------------------
# Example 1: Using an existing (built-in) style
# ---------------------------------------------
paragraph_1 = document.add_paragraph("This is the Normal Text")     # uses 'Normal' style
page1_heading1 = document.add_heading("This is Heading 1", level=1)
page1_heading2 = document.add_heading("This is Heading 2", level=2)
document.add_page_break()



# ------------------------------------------------------
# Example 2: Applying a style by name (important lesson)
# ------------------------------------------------------
paragraph_2 = document.add_paragraph("This is a title paragraph")
paragraph_2.style = "Title"
document.add_page_break()



# # --------------------------------------------------------------
# # Example 3: Trying to use a NON-EXISTING style (common mistake)
# # --------------------------------------------------------------
# paragraph_3 = document.add_paragraph("This text will NOT change")
# paragraph_3.style = "MyCustomStyle"   # Does Not exist
# document.add_page_break()



# ------------------------------------------------
# Example 4: Creating a custom style (CORRECT way)
# ------------------------------------------------

# 1️⃣ CREATE STYLE FIRST
style = document.styles.add_style("MyCustomStyle", WD_STYLE_TYPE.PARAGRAPH)
style.font.name = "Calibri"
style.font.size = Pt(16)
style.font.bold = True

# 2️⃣ NOW USE THE STYLE
paragraph_4 = document.add_paragraph("This text uses MyCustomStyle")
paragraph_4.style = "MyCustomStyle"
document.add_page_break()


# -----------------------------------------------------
# Example 5: Style inheritance (VERY important concept)
# -----------------------------------------------------
base_style = document.styles["Normal"]

new_style = document.styles.add_style("HighlightText", WD_STYLE_TYPE.PARAGRAPH)
new_style.base_style = base_style
new_style.font.size = Pt(14)
new_style.font.italic = True

paragraph_5 = document.add_paragraph("This is the Style Example paragraph", style="HighlightText")
document.add_page_break()



# ---------------------------------------------
# Example 6: Paragraph style vs Character style
# ---------------------------------------------
char_style = document.styles.add_style("RedText", WD_STYLE_TYPE.CHARACTER)
char_style.font.color.rbg = (255, 0, 0)
# char_style.font.bold = True

paragraph_6 = document.add_paragraph("Normal Text")
para6_add_run = paragraph_6.add_run(" RED TEXT")
para6_add_run.style = "RedText"






document.save("./ocr_data_doc/understanding_styles.docx")
print("Programm RUN Properlly..! || OK")
