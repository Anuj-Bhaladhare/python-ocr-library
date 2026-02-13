from docx import Document
from docx.shared import Inches
from docx.shared import Pt

document = Document()

# -------------------------------------------------
# Aided Heading 
# -------------------------------------------------
heading_para = document.add_paragraph()
heading_para_run = heading_para.add_run("Design and build \naccessible PDF tables")

heading_para_run.font.size = Pt(30)
heading_para_run.font.bold = True

heading_para_run_2 = heading_para.add_run("\nSample tables")
heading_para_run_2.font.size = Pt(20)
heading_para_run_2.font.bold = True


# -------------------------------------------------
# Table - 1
# -------------------------------------------------
table_heading = document.add_paragraph()
table_heading_run = table_heading.add_run("\n\nTable 1")
table_heading_run.font.size = Pt(15)
table_heading_run.font.bold = True

table_1_table = document.add_table(rows=3, cols=3)
first_row = table_1_table.rows[0]
second_row = table_1_table.rows[1]
third_row = table_1_table.rows[2]

first_row.cells[0].text = "Column header (TH)"
first_row.cells[1].text = "Column header (TH)"
first_row.cells[2].text = "Column header (TH)"

second_row.cells[0].text = "Row header (TH)"
second_row.cells[1].text = "Data cell (TD)"
second_row.cells[2].text = "Data cell (TD)"

third_row.cells[0].text = "Row header (TH)"
third_row.cells[1].text = "Data cell (TD)"
third_row.cells[2].text = "Data cell (TD)"

# -------------------------------------------------
# Table - 2
# -------------------------------------------------
table_2_heading = document.add_paragraph()
table_2_heading_run = table_2_heading.add_run("Table 2: example of footnotes referenced from within a table")
table_2_table = document.add_table(rows=8, cols=4)

# get cells
cell_1 = table_2_table.cell(0, 0)
cell_2 = table_2_table.cell(0, 1) 
cell_5 = table_2_table.cell(1, 0) 
cell_9 = table_2_table.cell(2, 0) 
cell_13 = table_2_table.cell(3, 0) 
cell_17 = table_2_table.cell(4, 0) 
cell_21 = table_2_table.cell(5, 0) 
cell_25 = table_2_table.cell(6, 0) 
cell_29 = table_2_table.cell(7, 0) 

# Cell Murging
cell_1_and_2 = cell_1.merge(cell_2)
# cell_5_and_9_and_13 = 
# cell_17_21_25_29 = 

document.save("./output_data/table-practice.docx")
print("Programm RUN Successfully | RUN OK..!")
