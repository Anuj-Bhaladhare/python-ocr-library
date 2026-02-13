from docx import Document

# # Opening a document
# document = Document()
# document.save('./ocr_data_doc/test.docx')

# # REALLY opening a document
# document = Document('./ocr_data_doc/test.docx')
# document.save('./ocr_data_doc/test.docx')

# Opening a ‘file-like’ document
f = open('./ocr_data_doc/test.docx', 'rb')
document = Document(f)
f.close()

# OR
with open('./ocr_data_doc/test.docx', 'rb') as f:
    source_stream = StringIO(f.read())
document = Document(source_stream)
source_stream.close()
...
target_stream = StringIO()
document.save(target_stream)








print("Programm RUN SUCCESS...!")



