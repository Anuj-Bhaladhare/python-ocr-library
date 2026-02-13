# Imports the library
from docx import Document
from docx.shared import Pt, Inches, Length, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE


# Create a document's
document = Document()

# -------------------
# Block-level vs. inline text objects
# -------------------
"""
Block-level elements start on a new line, take up the full width of their container (like <div>, <p>, <h1>), and create a "block" of content, while inline elements flow within the text, only taking up necessary width (like <span>, <a>, <strong>), and don't break the flow of text
"""
paragraph = document.add_paragraph(
    "This is a long paragraph that will automatically wrap to the next line "
    "when it reaches the page margin."
)

table = document.add_table(rows=1, cols=2)
cell_1 = table.cell(0, 0)
cell_2 = table.cell(0, 1)

cell_1.text = (
    "A"
)

cell_2.text = (
    "B" 
)

paragraph_2 = document.add_paragraph(
    "this is the second paragraph"
)

document.add_page_break()

# -----------------------
# Paragraph properties
# -----------------------
paragraph_3 = document.add_paragraph("This is BAD WAY to Apply 'Paragraph properties', It is NOT Adoptable for as Much as Industry Standerd")
paragraph_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER   # Text Center
paragraph_3.paragraph_format.space_before = Pt(12)                   # Space from Top like a Margin
paragraph_3.paragraph_format.space_after = Pt(12)


style = document.styles.add_style("MyParagraph", WD_STYLE_TYPE.PARAGRAPH)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER   # Text Center
style.paragraph_format.space_before = Pt(12)                   # Space from Top like a Margin
style.paragraph_format.space_after = Pt(12)

paragraph_4 = document.add_paragraph(
    "This is RIGHT WAY to Apply 'Paragraph properties', It is Industry Standerd",
    style="MyParagraph"
)


document.add_page_break()


# --------------------------
# Horizontal alignment (justification)
# --------------------------
paragraph_5 = document.add_paragraph("This is the paragraph is center aligned OK!")
paragraph_format_5 = paragraph_5.paragraph_format

# indicating alignment is inherited from the style hierarchy
# paragraph_format_5.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Right-side indent works in a similar way:
paragraph_format_5.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_format_5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
paragraph_format_5.alignment = WD_ALIGN_PARAGRAPH.LEFT


document.add_page_break()


indent_paragraph = document.add_paragraph(
    "Python is the programming language. "
    "It is widely used in data science, automation, and web development"
)
fmt = indent_paragraph.paragraph_format
# fmt.left_indent = Inches(0.5)
fmt.first_line_indent = Inches(0.25)

# Right-side indent works in a similar way:
# indent_paragraph

# Tab stops
tab_stop_paragraph = document.add_paragraph(
    "Name\tScore\n"
    "Anuj\t95\n"
    "Rahul\t88"
)

# Access tab stops via paragraph_format
tab_stops = tab_stop_paragraph.paragraph_format.tab_stops

# Add a tab stop at 0.5 inches
tab_stops.add_tab_stop(
    Inches(2),
    alignment=WD_TAB_ALIGNMENT.LEFT,
    leader=WD_TAB_LEADER.SPACES
)

# --------------------
# Line spacing
# --------------------
document.add_page_break()

line_spacing_paragraph = document.add_paragraph(
    "Line spacing is the distance between subsequent baselines in the lines of a paragraph. "
    "Line spacing can be specified either as an absolute distance or relative to the line height. "
    "This paragraph demonstrates double line spacing."
)
line_spacing_paragraph_format = line_spacing_paragraph.paragraph_format
line_spacing_paragraph_format.line_spacing = 2.0         # DOUBLE spacing
line_spacing_paragraph_format.space_after = Pt(12)       # optional: gap after paragraph
  
  




# ---------------------------
# Pagination properties
# ---------------------------
document.add_page_break()

# Chapter heading
heading = document.add_paragraph("Chapter 1: Introduction")
paragraph_format_heading = heading.paragraph_format

paragraph_format_heading.page_break_before = True       # start chapter on new page 
paragraph_format_heading.keep_with_next = True          # keep heading with content

# chapter content
content = document.add_paragraph(
    "This is the first paragraph of the chapter. "
    "It explains the topic in detail and should not be "
    "split awkwardly across pages."
)

paragraph_format_content = content.paragraph_format
paragraph_format_content.keep_together = True         # keep paragraph on one page
paragraph_format_content.widow_control = True         # avoid lonely lines

# -----------------------------
# Apply character formatting
# -----------------------------
character_formatting_paragraph = document.add_paragraph()

run1 = character_formatting_paragraph.add_run("Hello ")
run1.font.name = "Calibri"
run1.font.size = Pt(12)

run2 = character_formatting_paragraph.add_run("World \n")
run2.font.bold = True
run2.font.size = Pt(14)

run3 = character_formatting_paragraph.add_run("This is Italic FONT \n")
run3.font.italic = True
run3.font.size = Pt(12)

run4 = character_formatting_paragraph.add_run("This is UNDERLINE \n")
run4.font.underline = True
run4.font.size = Pt(12)

run5 = character_formatting_paragraph.add_run("This is the RBG Color")
run5.font.color.rgb = RGBColor(0x55, 0x80, 0x80)
run5.font.bold = True
run5.font.size = Pt(14)


document.save("./ocr_data_doc/working_with_text.docx")
# conformation code 
print("Programm Run Sucessfully...!")
